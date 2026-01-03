#!/usr/bin/env python3
"""
Generate sample files of multiple types for manual testing.
"""

from __future__ import annotations

import argparse
import random
from pathlib import Path
import zipfile

from docx import Document
from PIL import Image
from pypdf import PdfWriter


def _random_hex(length: int = 8) -> str:
	return "".join(random.choice("0123456789abcdef") for _ in range(length))


def _write_text(path: Path) -> None:
	stories = [
		"Once upon a time, a fox found a lantern in the woods.",
		"The lantern whispered a map to a hidden orchard.",
		"In the orchard, the fox met a watchmaker who kept time in jars.",
		"They agreed to trade a story for a seed, and the forest brightened.",
	]
	path.write_text(" ".join(random.sample(stories, k=3)), encoding="utf-8")


def _write_rtf(path: Path) -> None:
	text = "Once upon a time, a small boat sailed through a sky of stars."
	content = r"{\rtf1\ansi\deff0 " + text + "}"
	path.write_text(content, encoding="utf-8")


def _write_csv(path: Path) -> None:
	lines = ["name,score", "Ava,98", "Liam,91", "Noah,88"]
	path.write_text("\n".join(lines), encoding="utf-8")


def _write_tsv(path: Path) -> None:
	lines = ["name\tscore", "Ava\t98", "Liam\t91", "Noah\t88"]
	path.write_text("\n".join(lines), encoding="utf-8")


def _write_markdown(path: Path) -> None:
	content = "# Sample Notes\n\n- Idea: build a clockwork kite\n- Note: capture the wind\n"
	path.write_text(content, encoding="utf-8")


def _write_python(path: Path) -> None:
	content = "def greet(name):\n\treturn f\"Hello, {name}!\"\n"
	path.write_text(content, encoding="utf-8")


def _write_docx(path: Path) -> None:
	doc = Document()
	doc.add_heading("Sample Doc", level=1)
	doc.add_paragraph("A short tale about a clockwork bird.")
	doc.save(path)


def _write_pdf(path: Path) -> None:
	writer = PdfWriter()
	writer.add_blank_page(width=612, height=792)
	with path.open("wb") as handle:
		writer.write(handle)


def _write_solid_image(path: Path) -> None:
	color = tuple(random.randint(0, 255) for _ in range(3))
	image = Image.new("RGB", (240, 240), color=color)
	image.save(path)


def _write_gradient_image(path: Path) -> None:
	width, height = 240, 240
	image = Image.new("RGB", (width, height))
	pixels = image.load()
	for x in range(width):
		for y in range(height):
			pixels[x, y] = (x % 256, y % 256, (x + y) % 256)
	image.save(path)


def _write_epub(path: Path) -> None:
	mimetype = "application/epub+zip"
	container_xml = """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>
"""
	opf_xml = """<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="BookId" version="2.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:title>Sample EPUB Guide</dc:title>
    <dc:creator>Test Author</dc:creator>
    <dc:subject>Documentation</dc:subject>
    <dc:description>A short sample EPUB file.</dc:description>
  </metadata>
  <manifest>
    <item id="chap1" href="chapter1.xhtml" media-type="application/xhtml+xml"/>
  </manifest>
  <spine>
    <itemref idref="chap1"/>
  </spine>
</package>
"""
	chapter_xml = """<?xml version="1.0" encoding="utf-8"?>
<html xmlns="http://www.w3.org/1999/xhtml">
  <head><title>Sample EPUB Guide</title></head>
  <body><p>This is a sample EPUB chapter.</p></body>
</html>
"""
	with zipfile.ZipFile(path, "w") as archive:
		archive.writestr("mimetype", mimetype, compress_type=zipfile.ZIP_STORED)
		archive.writestr("META-INF/container.xml", container_xml)
		archive.writestr("OEBPS/content.opf", opf_xml)
		archive.writestr("OEBPS/chapter1.xhtml", chapter_xml)


def generate_samples(output_dir: Path) -> None:
	output_dir.mkdir(parents=True, exist_ok=True)
	_write_text(output_dir / f"{_random_hex()}_story.txt")
	_write_rtf(output_dir / f"{_random_hex()}_story.rtf")
	_write_markdown(output_dir / f"{_random_hex()}_notes.md")
	_write_csv(output_dir / f"{_random_hex()}_data.csv")
	_write_tsv(output_dir / f"{_random_hex()}_data.tsv")
	_write_python(output_dir / f"{_random_hex()}_script.py")
	_write_docx(output_dir / f"{_random_hex()}_letter.docx")
	_write_pdf(output_dir / f"{_random_hex()}_sheet.pdf")
	_write_solid_image(output_dir / f"{_random_hex()}_solid.png")
	_write_gradient_image(output_dir / f"{_random_hex()}_gradient.jpg")
	_write_epub(output_dir / f"{_random_hex()}_sample.epub")


def main() -> None:
	parser = argparse.ArgumentParser(
		description="Generate sample files of multiple types for testing."
	)
	parser.add_argument(
		"-o",
		"--output",
		dest="output",
		default="tests/generated_samples",
		help="Output directory for generated sample files.",
	)
	args = parser.parse_args()
	generate_samples(Path(args.output))


if __name__ == "__main__":
	main()
